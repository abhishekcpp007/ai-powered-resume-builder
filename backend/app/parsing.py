import io
import re
from typing import List, Tuple

import pdfplumber
from docx import Document


class ParseError(ValueError):
    pass

_SUMMARY_HEADINGS = ("summary", "objective", "profile", "about")
_SKILLS_HEADINGS = ("skills", "technical skills", "core competencies", "technologies")
_ANY_HEADING_RE = re.compile(r"^[A-Za-z][A-Za-z /&]{2,30}$")


def _find_section(lines: List[str], headings: Tuple[str, ...]) -> str:
    for i, line in enumerate(lines):
        stripped = line.strip().strip(":")
        if stripped.lower() in headings:
            section_lines = []
            for candidate in lines[i + 1:]:
                candidate_stripped = candidate.strip().strip(":")
                if not candidate_stripped:
                    if section_lines:
                        break
                    continue
                if (
                    _ANY_HEADING_RE.match(candidate_stripped)
                    and candidate_stripped.lower() not in headings
                    and (candidate_stripped.isupper() or len(candidate_stripped.split()) <= 4)
                    and len(section_lines) > 0
                ):
                    break
                section_lines.append(candidate_stripped)
            return " ".join(section_lines).strip()
    return ""


def extract_summary_and_skills(resume_text: str) -> Tuple[str, List[str]]:
    lines = resume_text.splitlines()

    summary = _find_section(lines, _SUMMARY_HEADINGS)

    skills_text = _find_section(lines, _SKILLS_HEADINGS)
    skills: List[str] = []
    if skills_text:
        parts = re.split(r"[,•|/]|\s{2,}", skills_text)
        skills = [p.strip() for p in parts if p.strip()]

    return summary, skills


def extract_text(filename: str, data: bytes) -> str:
    lower = filename.lower()
    if lower.endswith(".pdf"):
        return _extract_pdf(data)
    if lower.endswith(".docx"):
        return _extract_docx(data)
    raise ParseError("Unsupported file type. Only .pdf and .docx are supported.")


def _extract_pdf(data: bytes) -> str:
    try:
        with pdfplumber.open(io.BytesIO(data)) as pdf:
            pages = [page.extract_text() or "" for page in pdf.pages]
        return "\n".join(pages).strip()
    except Exception as e:
        raise ParseError(f"Failed to read PDF: {e}") from e


def _extract_docx(data: bytes) -> str:
    try:
        document = Document(io.BytesIO(data))
        parts = [p.text for p in document.paragraphs]
        for table in document.tables:
            for row in table.rows:
                for cell in row.cells:
                    parts.append(cell.text)
        return "\n".join(p for p in parts if p.strip()).strip()
    except Exception as e:
        raise ParseError(f"Failed to read DOCX: {e}") from e
