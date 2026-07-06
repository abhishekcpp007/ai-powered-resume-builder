import io
from xml.sax.saxutils import escape as _xml_escape

from docx import Document
from docx.shared import Pt
from reportlab.lib.pagesizes import LETTER
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import inch
from reportlab.platypus import ListFlowable, ListItem, Paragraph, SimpleDocTemplate, Spacer

from .schemas import TailoredResume


def resume_to_text(resume: TailoredResume) -> str:
    lines = []
    if resume.name:
        lines.append(resume.name)
    if resume.contact:
        lines.append(resume.contact)
    if resume.name or resume.contact:
        lines.append("")

    if resume.summary:
        lines.append("SUMMARY")
        lines.append(resume.summary)
        lines.append("")

    if resume.skills:
        lines.append("SKILLS")
        lines.append(", ".join(resume.skills))
        lines.append("")

    if resume.experience:
        lines.append("EXPERIENCE")
        for job in resume.experience:
            header = " — ".join(p for p in [job.title, job.company] if p)
            meta = " | ".join(p for p in [job.location, job.dates] if p)
            lines.append(header + (f" ({meta})" if meta else ""))
            for bullet in job.bullets:
                lines.append(f"  - {bullet}")
            lines.append("")

    if resume.education:
        lines.append("EDUCATION")
        for edu in resume.education:
            header = " — ".join(p for p in [edu.degree, edu.school] if p)
            lines.append(header + (f" ({edu.dates})" if edu.dates else ""))
        lines.append("")

    for section in resume.other_sections:
        if section.heading:
            lines.append(section.heading.upper())
        if section.content:
            lines.append(section.content)
        lines.append("")

    return "\n".join(lines).strip() + "\n"


def resume_to_docx(resume: TailoredResume) -> bytes:
    document = Document()

    if resume.name:
        document.add_heading(resume.name, level=1)
    if resume.contact:
        p = document.add_paragraph(resume.contact)
        p.runs[0].font.size = Pt(10)

    def add_section_heading(text: str) -> None:
        document.add_heading(text, level=2)

    if resume.summary:
        add_section_heading("Summary")
        document.add_paragraph(resume.summary)

    if resume.skills:
        add_section_heading("Skills")
        document.add_paragraph(", ".join(resume.skills))

    if resume.experience:
        add_section_heading("Experience")
        for job in resume.experience:
            header = " — ".join(p for p in [job.title, job.company] if p)
            meta = " | ".join(p for p in [job.location, job.dates] if p)
            p = document.add_paragraph()
            run = p.add_run(header)
            run.bold = True
            if meta:
                p.add_run(f"  ({meta})")
            for bullet in job.bullets:
                document.add_paragraph(bullet, style="List Bullet")

    if resume.education:
        add_section_heading("Education")
        for edu in resume.education:
            header = " — ".join(p for p in [edu.degree, edu.school] if p)
            text = header + (f" ({edu.dates})" if edu.dates else "")
            document.add_paragraph(text)

    for section in resume.other_sections:
        if section.heading:
            add_section_heading(section.heading)
        if section.content:
            document.add_paragraph(section.content)

    buffer = io.BytesIO()
    document.save(buffer)
    return buffer.getvalue()


def resume_to_pdf(resume: TailoredResume) -> bytes:
    buffer = io.BytesIO()
    doc = SimpleDocTemplate(
        buffer,
        pagesize=LETTER,
        topMargin=0.6 * inch,
        bottomMargin=0.6 * inch,
        leftMargin=0.7 * inch,
        rightMargin=0.7 * inch,
    )

    styles = getSampleStyleSheet()
    name_style = ParagraphStyle("NameStyle", parent=styles["Title"], fontSize=18, spaceAfter=2)
    contact_style = ParagraphStyle("ContactStyle", parent=styles["Normal"], fontSize=9, spaceAfter=10)
    heading_style = ParagraphStyle(
        "SectionHeading", parent=styles["Heading2"], fontSize=12, spaceBefore=10, spaceAfter=4,
        textColor="#1a1a1a",
    )
    body_style = ParagraphStyle("BodyStyle", parent=styles["Normal"], fontSize=10, leading=14)
    job_header_style = ParagraphStyle("JobHeader", parent=styles["Normal"], fontSize=10.5, leading=14, spaceBefore=6)

    esc = _xml_escape

    story = []
    if resume.name:
        story.append(Paragraph(esc(resume.name), name_style))
    if resume.contact:
        story.append(Paragraph(esc(resume.contact), contact_style))

    if resume.summary:
        story.append(Paragraph("Summary", heading_style))
        story.append(Paragraph(esc(resume.summary), body_style))

    if resume.skills:
        story.append(Paragraph("Skills", heading_style))
        story.append(Paragraph(esc(", ".join(resume.skills)), body_style))

    if resume.experience:
        story.append(Paragraph("Experience", heading_style))
        for job in resume.experience:
            header = " — ".join(p for p in [job.title, job.company] if p)
            meta = " | ".join(p for p in [job.location, job.dates] if p)
            text = f"<b>{esc(header)}</b>" + (f"  ({esc(meta)})" if meta else "")
            story.append(Paragraph(text, job_header_style))
            if job.bullets:
                items = [ListItem(Paragraph(esc(b), body_style)) for b in job.bullets]
                story.append(ListFlowable(items, bulletType="bullet", leftIndent=16))

    if resume.education:
        story.append(Paragraph("Education", heading_style))
        for edu in resume.education:
            header = " — ".join(p for p in [edu.degree, edu.school] if p)
            text = header + (f" ({edu.dates})" if edu.dates else "")
            story.append(Paragraph(esc(text), body_style))

    for section in resume.other_sections:
        if section.heading:
            story.append(Paragraph(esc(section.heading), heading_style))
        if section.content:
            for line in section.content.splitlines():
                if line.strip():
                    story.append(Paragraph(esc(line), body_style))

    if not story:
        story.append(Paragraph("Empty resume", body_style))

    doc.build(story)
    return buffer.getvalue()
